"""Generate Word document from main.tex content."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Styles ───────────────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name="Times New Roman", font_size=11, bold=False, space_before=6, space_after=3):
    try:
        s = styles[style_name]
    except KeyError:
        return
    s.font.name = font_name
    s.font.size = Pt(font_size)
    s.font.bold = bold
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after  = Pt(space_after)

set_style("Normal", font_size=11, space_before=0, space_after=4)
set_style("Heading 1", font_size=13, bold=True, space_before=10, space_after=4)
set_style("Heading 2", font_size=11, bold=True, space_before=8, space_after=2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def clean(text):
    """Strip LaTeX markup, returning plain unicode."""
    # Remove comments
    text = re.sub(r'%.*', '', text)
    # \emph{x} -> x (italic marker kept for runs)
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    # \textit{x} -> x
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    # \textbf{x} -> x
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    # \mathbf{x} -> x
    text = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', text)
    # \bfseries
    text = text.replace(r'\bfseries', '')
    # Math mode: strip $ ... $
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    # \cite{...} -> [?]
    text = re.sub(r'\\cite\{[^}]*\}', '[?]', text)
    # \Cref{...} -> Figure/Table ref
    text = re.sub(r'\\Cref\{tab:([^}]*)\}', r'Table', text)
    text = re.sub(r'\\Cref\{fig:([^}]*)\}', r'Figure', text)
    text = re.sub(r'\\ref\{[^}]*\}', '', text)
    # \url{...}
    text = re.sub(r'\\url\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    # Spacing commands
    for cmd in [r'\\,', r'\\;', r'\\:', r'\\!', r'\\ ', r'\\quad', r'\\qquad']:
        text = text.replace(cmd, ' ')
    # Common commands
    text = text.replace(r'\%', '%').replace(r'\$', '$').replace(r'\&', '&')
    text = text.replace(r'\_', '_').replace(r'\^', '^').replace(r'\#', '#')
    text = text.replace(r'\{', '{').replace(r'\}', '}')
    text = text.replace(r'\emph', '').replace(r'\textbf', '').replace(r'\textit', '')
    text = text.replace(r'\\', ' ').replace(r'\to', '→')
    text = text.replace('--', '–').replace('---', '—')
    text = text.replace(r'\times', '×').replace(r'\approx', '≈')
    text = text.replace(r'\pm', '±').replace(r'\geq', '≥').replace(r'\leq', '≤')
    text = text.replace(r'\alpha', 'α').replace(r'\beta', 'β').replace(r'\gamma', 'γ')
    text = text.replace(r'\tau', 'τ').replace(r'\sigma', 'σ').replace(r'\mu', 'μ')
    text = text.replace(r'\in', '∈').replace(r'\ldots', '…').replace(r'\cdots', '…')
    text = text.replace(r'\hat{a}', 'â').replace(r'\mathbf{u}', 'u')
    text = text.replace(r'\sim', '~').replace(r'\infty', '∞')
    text = text.replace('{,}', ',').replace(r'\,', ' ')
    text = re.sub(r'\{([^}]*)\}', r'\1', text)  # remaining braces
    text = re.sub(r'\\[a-zA-Z]+\*?', '', text)  # remaining commands
    text = re.sub(r' +', ' ', text)
    return text.strip()

def add_heading(text, level=1):
    p = doc.add_heading(clean(text), level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(text, indent=False, italic_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if italic_prefix:
        run = p.add_run(italic_prefix + " ")
        run.italic = True
        run.bold   = True
    p.add_run(clean(text))
    return p

def add_table_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(clean(text))
    run.font.size = Pt(9)
    run.italic = True

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE & AUTHORS
# ═══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Urban Expansion Monitoring Using Transfer Learning on Historical Satellite Imagery")
title_run.bold = True
title_run.font.size = Pt(16)

doc.add_paragraph()

authors_p = doc.add_paragraph()
authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors = [
    "Kaustubh Agarwal", "Prajwal Gupta", "Kabir Rai",
    "P Manideep Reddy", "Naman Khurana", "Dr. Divyashree N"
]
authors_p.add_run(", ".join(authors)).font.size = Pt(11)

affil_p = doc.add_paragraph()
affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil_p.add_run("Department of Computer Science and Engineering, PES University, Bengaluru, India").font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Abstract", level=1)
abstract = (
    "Uncontrolled urban expansion in Indian metropolitan regions accelerates encroachment on legally protected zones, "
    "yet existing studies rely on single-city, single-seed SVM/RF benchmarks without downstream forecasting or alerting. "
    "We propose an end-to-end framework integrating transfer-learning classification, SAR–optical fusion, "
    "self-supervised pre-training, uncertainty-aware Bi-LSTM forecasting, and regulatory encroachment alerting. "
    "A benchmark of 2,730 Sentinel-2 patches across Mumbai, Delhi NCR, and Bangalore with a novel three-class taxonomy "
    "(Urban, Non-Urban, Transition) is constructed from ESA WorldCover 2021. Six model families (SVM, Random Forest, "
    "MobileNetV3-Small, EfficientNet-B0, Swin-Tiny, ResNet50) are trained with three-stage progressive fine-tuning and "
    "a combined CE–focal–Dice loss across three random seeds. ResNet50 achieves 97.5 ± 0.2% overall accuracy (OA), "
    "surpassing the best published Indian SVM (91.01%) by six points. An improved SVM reaches 92.6 ± 1.4%, beating the "
    "published baseline but trailing deep learning by five points. The first Leave-One-City-Out (LOCO) benchmark for "
    "Indian metros reveals a 15–20% cross-city accuracy drop and a CNN–Transformer ranking reversal where Swin-Tiny "
    "(79.1 ± 3.9%) generalizes better than ResNet50 (77.1 ± 5.1%). The pipeline produces uncertainty-calibrated "
    "2024–2035 forecasts and, in a seven-city encroachment simulation, generates 55 simulated alerts with rule-based "
    "routing to the appropriate Indian regulatory authority."
)
add_para(abstract)

kw_p = doc.add_paragraph()
kw_run = kw_p.add_run("Keywords: ")
kw_run.bold = True
kw_run.italic = True
kw_p.add_run(
    "Urban expansion monitoring, transfer learning, Sentinel-2, SAR–optical fusion, self-supervised learning, "
    "Bi-LSTM forecasting, cross-city generalization, Indian metropolitan regions, regulatory encroachment alerts."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  I. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("I. Introduction", level=1)
add_heading("A. Problem Domain", level=2)
add_para(
    "India's urban population is projected to exceed 600 million by 2031, with major metropolitan regions already "
    "expanded two- to four-fold in built-up area since 1990. This growth triggers three interlinked problems that "
    "ground inspection cannot solve. First, encroachment on legally protected zones continues despite the Coastal "
    "Regulation Zone Notification 2019, the Forest Conservation Act 1980, and the Wetland Rules 2017; Mumbai's Sanjay "
    "Gandhi National Park, Bangalore's Bellandur wetland, and Delhi's Yamuna floodplain are routinely breached by "
    "informal construction. Second, infrastructure planning ignores uncertainty — point forecasts offer no confidence "
    "bounds, causing capital misallocation in Indian smart-city investments. Third, deep learning is under-used for "
    "Indian urban classification: backbones such as ResNet50, EfficientNet, Swin, and MobileNetV3 routinely exceed 97% "
    "on European benchmarks, yet published Indian urban studies rely overwhelmingly on single-city SVM/RF pipelines "
    "reaching only 89–91% OA, with no cross-city evaluation and no link from classification to forecasting or alerting. "
    "Cloud-masked Sentinel-2/1 and harmonized Landsat composites via Google Earth Engine (GEE) provide the coverage; "
    "what is missing is an integrated pipeline from raw Indian imagery to policy-actionable output."
)

add_heading("B. Key Contributions", level=2)
contributions = [
    ("Multi-city Indian urban expansion benchmark with the first LOCO evaluation for Indian metros.",
     "We construct 2,730 Sentinel-2 patches across Mumbai, Delhi NCR, and Bangalore (2017–2023), automatically "
     "labeled from ESA WorldCover 2021 with three classes (Urban, Non-Urban, Transition), and evaluate six models "
     "with three-seed statistical rigor (seeds 42, 123, 7)."),
    ("Progressive fine-tuning achieving 97.5 ± 0.2% OA",
     "on real Indian data without any domain-specific pre-training, exceeding the best published Indian SVM benchmark "
     "by more than six points."),
    ("Novel Transition class via morphological dilation",
     "of urban boundaries (100 m buffer), which models mixed-pixel ambiguity at the active urban–rural interface "
     "and is the class most relevant to encroachment detection."),
    ("Uncertainty-aware Bi-LSTM forecasting",
     "producing 95% Monte-Carlo-Dropout confidence intervals on 2024–2035 urban-area forecasts — a capability absent "
     "from every published Indian urban forecasting study we surveyed."),
    ("First weak-supervision-to-policy-action pipeline for Indian urban monitoring: classify → time series → forecast → regulatory alert.",
     "The system encodes ten India-specific regulatory zone types and 30+ named protected areas, with alerts routed "
     "to MoEFCC, CZMA, or the State Forest Department."),
    ("CNN–Transformer ranking reversal under cross-city transfer.",
     "ResNet50 wins in-distribution, but Swin-Tiny generalizes better under LOCO (79.1% vs. 77.1% OA), empirically "
     "confirming the OOD-robustness hypothesis of Bai et al. on Indian satellite data."),
    ("Transparent negative-result reporting",
     "for SAR fusion, SimCLR pre-training, and FPN ablation, each with three-seed statistical bounds — addressing "
     "positive-result bias in remote sensing."),
]
for i, (title, detail) in enumerate(contributions, 1):
    p = doc.add_paragraph(style='List Number')
    r1 = p.add_run(title + " ")
    r1.bold = True
    r1.italic = True
    p.add_run(detail)

# ═══════════════════════════════════════════════════════════════════════════════
#  II. LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("II. Literature Survey", level=1)
add_para(
    "Table 1 summarizes sixteen related works spanning backbones, Indian urban classification, cross-city transfer, "
    "change detection, and forecasting. The discussion below identifies the gap we fill."
)

# Literature survey table
lit_caption = doc.add_paragraph()
lit_caption.add_run("Table 1: Literature Survey").bold = True
lit_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

lit_table = doc.add_table(rows=1, cols=6)
lit_table.style = 'Table Grid'
hdr = lit_table.rows[0].cells
for cell, txt in zip(hdr, ["Author(s)", "Dataset", "Year", "Methodology", "Result", "Comments"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True

lit_rows = [
    ("He et al.", "ImageNet", "2016", "ResNet50 backbone", "96.78% (EuroSAT)", "Foundation backbone"),
    ("Tan & Le", "ImageNet", "2019", "EfficientNet compound NAS", "97.1% (EuroSAT)", "Accuracy/FLOPs trade-off"),
    ("Helber et al.", "EuroSAT", "2019", "CNN + 13 S2 bands", "98.57% OA", "Standard RS benchmark"),
    ("Liu et al.", "ImageNet", "2021", "Shifted-window attention", "SOTA on RS", "Best LOCO generalizer (ours)"),
    ("Howard et al.", "ImageNet", "2019", "NAS lightweight CNN", "Lightweight CNN", "91.5% OA (ours)"),
    ("Chamoli et al.", "S2 India", "2024", "SVM, RF", "91.01%; 89.67%", "Best Indian SVM"),
    ("Katpadi et al.", "S2 India", "2025", "IRv2+UNet ensemble", "98.21% OA", "Binary, single region"),
    ("Sharma & Joshi", "S1+S2 Delhi", "2022", "Fusion CNN", "92.0% OA", "SAR helps under cloud"),
    ("Bai et al.", "ImageNet var.", "2021", "CNN vs. ViT OOD", "ViT better OOD", "Confirmed on our data"),
    ("Li et al.", "C2Seg (BE,BJ)", "2023", "Domain adaptation", "SOTA C2Seg", "First Indian LOCO is ours"),
    ("Wang et al.", "Chinese cities", "2023", "Deep transfer", "SOTA Chinese", "Chinese only"),
    ("Guo et al.", "LEVIR-CD", "2023", "Global-aware Siamese", "F1 = 91.21%", "Current CD SOTA"),
    ("Singh et al.", "Sat+census", "2025", "CA-Markov + ML", "OA = 93.6%", "Point forecasts only"),
    ("Yadav & Agarwal", "Lucknow India", "2025", "Bi-LSTM + ANN", "Competitive MAE", "No CI bands"),
    ("Kalinicheva et al.", "89 S2 studies", "2025", "Systematic review", "15–25% deploy gap", "Validates our 18%"),
    ("World Bank", "Global sat.", "2024", "CNN + alerts", "General alerts", "No CI, no Indian law"),
]
for row_data in lit_rows:
    row = lit_table.add_row().cells
    for cell, val in zip(row, row_data):
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

lit_topics = [
    ("Transfer learning for EO.",
     "On EuroSAT, ResNet50, EfficientNet, and ViT reach 96.78–98.57% OA, but transferability to Indian imagery is unexplored."),
    ("Indian urban classification",
     "is dominated by SVM/RF (best: 91.01%); binary single-region ensembles reach 98.21% but none perform multi-city LOCO analysis."),
    ("Cross-city transfer",
     "and a 2025 review of 89 S2 studies report 15–25% deploy gaps; our 18% confirms this for Indian metros."),
    ("Siamese change detection",
     "on LEVIR-CD has reached F1 = 91.21% but does not forecast or alert."),
    ("CNN vs. transformer:",
     "transformers outperform CNNs OOD while CNNs win on small datasets — consistent with our reversal."),
    ("Indian forecasting",
     "relies on CA-Markov without uncertainty quantification."),
    ("The gap:",
     "no published work connects, for Indian metros, multi-seed DL classification, LOCO generalization, uncertainty-aware forecasting, and regulatory alerting in a single pipeline."),
]
for title, detail in lit_topics:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title + " ")
    r1.italic = True
    p.add_run(detail)

# ═══════════════════════════════════════════════════════════════════════════════
#  III. DATASET CONSTRUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("III. Dataset Construction", level=1)
add_heading("A. Study Area, Data Sources, and Processing", level=2)
add_para(
    "Three Indian metropolitan archetypes are selected: Mumbai (coastal megacity, constrained by the Arabian Sea and "
    "Sanjay Gandhi NP), Delhi NCR (radial sprawl with satellite towns), and Bangalore (IT-corridor city with dispersed "
    "built-up nuclei). All satellite imagery is exported via the Google Earth Engine (GEE) Python API (~5.7 GB total). "
    "Sentinel-2 L2A (10 m, 6 bands B2/3/4/8/11/12, 2017–2023, SCL cloud-masked, pre-monsoon composites) serves as "
    "the primary classification input; Sentinel-1 SAR (10 m, VV/VH, 2023 post-monsoon, 3×3 speckle-filtered, "
    "dB-clipped [−30, 0]) is used for the fusion experiment; Landsat 5/7/8/9 (30 m, harmonized 6-band composites for "
    "1990/2000/2010/2023) provides historical anchors for the Bi-LSTM forecaster only. From these GeoTIFFs, 256×256 "
    "patches are extracted at the native sensor resolution, producing 2,730 optical patches and 910 co-located SAR "
    "patches. The dataset is split 70/15/15 (train/val/test) with three additional LOCO folds for cross-city evaluation."
)

add_heading("B. Labels, Taxonomy, and Dataset Statistics", level=2)
label_items = [
    ("Labels —", "derived via weak supervision from ESA WorldCover 2021 (10 m, reported 74.4% global accuracy), cross-validated against Google Dynamic World."),
    ("Taxonomy —", "Three classes: Urban (WorldCover class 50), Non-Urban (all others), and Transition (100 m morphological dilation of Urban boundaries — novel for Indian urban studies, absorbing mixed-pixel ambiguity at the active expansion interface)."),
    ("Dataset Statistics —", "2,730 patches (Mumbai ≈900, Delhi NCR ≈1,200, Bangalore ≈630); class distribution: Urban 19.2%, Non-Urban 62.4%, Transition 18.4%. For change detection we additionally use LEVIR-CD (637 bi-temporal pairs, 0.5 m). Bi-LSTM features include Census population, GSDP, Smart City/AMRUT allocations, metro rail, NH density, SEZ/IT parks, green cover, and policy events."),
]
for title, detail in label_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title + " ")
    r1.italic = True
    p.add_run(detail)

# ═══════════════════════════════════════════════════════════════════════════════
#  IV. PROPOSED METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("IV. Proposed Methodology", level=1)
add_para(
    "The proposed end-to-end framework comprises a transfer-learning classifier, an FPN-enhanced feature head, "
    "SAR–optical fusion, SimCLR self-supervised pre-training, a Bi-LSTM temporal forecaster, and a regulatory-aware "
    "alert engine."
)

add_heading("A. Backbones, FPN Head, and Progressive Fine-Tuning", level=2)
backbone_items = [
    ("Backbones —", "Four deep backbones (all ImageNet-pretrained) are evaluated alongside SVM and Random Forest: ResNet50 (11.3M, primary), EfficientNet-B0 (5.3M, ablation backbone), Swin-Tiny (29.8M, best cross-city generalizer), and MobileNetV3-Small (3.4M, edge deployment)."),
    ("FPN Head —", "A three-level Feature Pyramid Network (FPN) at strides 4/8/16 is upsampled and concatenated before the classifier head."),
    ("Progressive Fine-Tuning —", "Training uses three-stage progressive fine-tuning with cosine LR annealing and early stopping (patience 10): Stage 1 trains the head only (5 epochs, lr 10⁻³), Stage 2 unfreezes the last two backbone blocks (lr 10⁻⁴), and Stage 3 fine-tunes the full network (lr 10⁻⁵)."),
]
for title, detail in backbone_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title + " ")
    r1.italic = True
    p.add_run(detail)

add_heading("B. Combined Loss", level=2)
add_para(
    "The training loss combines cross-entropy, focal, and Dice terms:\n\n"
    "    L = 0.6·L_CE + 0.3·L_Focal(γ=2) + 0.1·L_Dice\n\n"
    "Focal loss down-weights easy samples, Dice improves minority-class recall, and CE provides probabilistic "
    "calibration. Class weights [1.0, 1.0, 3.0] upweight the minority Transition class."
)

add_heading("C. SAR–Optical Fusion and SimCLR Pre-Training", level=2)
add_para(
    "A dual-branch encoder processes 6-channel optical and 2-channel SAR inputs separately; feature maps are "
    "concatenated before the shared FPN head. Training uses 910 aligned optical–SAR pairs (EfficientNet-B0) against "
    "an optical-only baseline. In parallel, SimCLR contrastive pre-training is run on unlabeled Indian patches "
    "(random crop, color jitter, Gaussian blur; projection dim 128; τ = 0.07; 100 epochs) and then fine-tuned with "
    "the progressive strategy."
)

add_heading("D. Bi-LSTM Forecaster and Alert Engine", level=2)
add_para(
    "For bi-temporal change detection on LEVIR-CD, a shared-weight Siamese EfficientNet-B0 encoder extracts features "
    "from each temporal image; the absolute feature difference is passed to a two-layer MLP classification head. "
    "The forecasting module is a bidirectional LSTM with residual connections, LayerNorm, 2-head temporal attention, "
    "and Monte-Carlo Dropout (p = 0.2) yielding 94K parameters. Each time step receives the satellite-derived urban "
    "area plus fifteen socio-economic and policy features. Temporal split: train 1990–2015, validation 2016–2019, "
    "test 2020–2023. Fifty MC forward passes produce 95% CIs on 2024–2035 forecasts. A three-head change detector "
    "classifies each patch for (i) change/no-change, (ii) severity ∈ {NONE, LOW, MED, HIGH, CRIT}, and "
    "(iii) encroachment type, encoding ten India-specific regulatory zone types (CRZ-I/II/III, Forest Reserve, "
    "Protected Forest, Wetland, Lake Buffer, River Floodplain, Green Belt, Western Ghats ESA) and 30+ named "
    "protected areas; alerts are auto-routed to MoEFCC, CZMA, or the State Forest Department."
)
add_para(
    "At inference time, the integrated pipeline slides a 256×256 window over each GeoTIFF, counts Urban pixels to "
    "build the city-year area series u_c = [a_{c,1990}, …, a_{c,2023}], runs fifty MC-Dropout forward passes through "
    "the Bi-LSTM to obtain â_{c,t} with 95% CI for t ∈ [2024, 2035], intersects each forecast zone with the "
    "regulatory database, and routes each alert to the correct authority."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  V. RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("V. Result Analysis and Discussions", level=1)
add_para(
    "All results are reported as mean ± std over three random seeds {42, 123, 7}; significance uses paired t-tests at p < 0.05."
)

add_heading("A. Main Benchmark on Real Indian Data", level=2)

# Table 2: Main benchmark
cap2 = doc.add_paragraph()
cap2.add_run("Table 2: Main benchmark on real Indian satellite data (3 seeds, mean ± std). "
             "†: improved with NDVI/NDBI/NDWI/SAVI/BSI + texture + PCA + grid search.").bold = True
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
for cell, txt in zip(t2.rows[0].cells, ["Model", "OA (%)", "F1", "mIoU"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True

main_rows = [
    ("SVM (raw pixels)",        "89.2 ± 0.4", "0.890 ± 0.005", "0.763 ± 0.009"),
    ("SVM (improved†)",         "92.6 ± 1.4", "0.899 ± 0.018", "0.825 ± 0.027"),
    ("Random Forest",           "88.2 ± 1.7", "0.875 ± 0.017", "0.733 ± 0.028"),
    ("RF (improved†)",          "90.7 ± 1.4", "0.871 ± 0.017", "0.784 ± 0.026"),
    ("MobileNetV3-Small",       "91.5 ± 1.6", "0.920 ± 0.014", "0.823 ± 0.028"),
    ("EfficientNet-B0",         "93.4 ± 2.3", "0.936 ± 0.022", "0.857 ± 0.044"),
    ("Swin-Tiny",               "93.6 ± 2.6", "0.939 ± 0.024", "0.862 ± 0.046"),
    ("ResNet50 (best)",         "97.5 ± 0.2", "0.976 ± 0.002", "0.939 ± 0.005"),
]
for i, row_data in enumerate(main_rows):
    row = t2.add_row().cells
    for cell, val in zip(row, row_data):
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if row_data[0] == "ResNet50 (best)":
            run.bold = True

doc.add_paragraph()
add_para(
    "ResNet50 wins on every metric and has the lowest seed variance, indicating both accuracy and stability, "
    "significantly outperforming MobileNetV3-Small (p = 0.038). Our raw-SVM result (89.2%) matches published Indian "
    "SVM baselines (89–91%), confirming that WorldCover 2021 labels — despite 74.4% global accuracy — are reliable "
    "for dense Indian metros; the improved-SVM (92.6 ± 1.4%) surpasses the best published Indian SVM — all three "
    "seeds exceed 91.01% — while trailing ResNet50 by ~5 points, establishing a hard ceiling for traditional ML. "
    "Our 97.5% OA is within 0.7 points of the highest published Indian result — IRUNet at 98.21% — despite being "
    "harder (3-class, multi-city, LOCO). ResNet50 Transition recall ranges 78.8–96.9% across cities; Bangalore's "
    "lower value reflects IT-park vegetation boundary ambiguity. No prior Indian study combines multi-seed DL "
    "classification, cross-city LOCO, uncertainty-aware forecasting, and regulatory alerting in one pipeline."
)

add_heading("B. Cross-City Generalization and CNN–Transformer Reversal", level=2)

cap3 = doc.add_paragraph()
cap3.add_run("Table 3: Leave-One-City-Out cross-city benchmark (3 seeds, mean ± std).").bold = True
cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER

t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
for cell, txt in zip(t3.rows[0].cells, ["Model", "OA (%)", "F1", "mIoU"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True

loco_rows = [
    ("EfficientNet-B0", "76.8 ± 2.2", "0.775 ± 0.011", "0.525 ± 0.006"),
    ("ResNet50",        "77.1 ± 5.1", "0.778 ± 0.037", "0.542 ± 0.026"),
    ("Swin-Tiny (best)","79.1 ± 3.9", "0.797 ± 0.034", "0.567 ± 0.034"),
]
for row_data in loco_rows:
    row = t3.add_row().cells
    for cell, val in zip(row, row_data):
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if "best" in row_data[0]:
            run.bold = True

doc.add_paragraph()
add_para(
    "Two findings stand out. First, a CNN–Transformer ranking reversal: ResNet50 drops 20.4% (97.5%→77.1%) under "
    "LOCO while Swin-Tiny drops only 14.5% (93.6%→79.1%), consistent with the OOD-robustness hypothesis of Bai et al. "
    "Self-attention learns abstract, transferable urban patterns while CNN filters overfit to city-specific textures "
    "(Mumbai's coastlines, Delhi's road grids). GradCAM analysis confirms this: ResNet50 activations concentrate on "
    "local texture patches, whereas Transition-class heatmaps show diffuse edge attention at construction fronts. "
    "Second, a 14.5–20.4% per-model in-distribution-to-LOCO gap independently replicates the 15–25% deploy gap "
    "reported in a 2025 systematic review of 89 Sentinel-2 studies. Delhi NCR is the hardest held-out target "
    "(71.8–76.7%) owing to its gradual radial sprawl, while Bangalore is the easiest (80.4–83.9%) despite being "
    "the hardest in-distribution city — a separation between local ambiguity and cross-city domain shift that is "
    "itself a novel finding."
)

add_heading("C. Ablation, SAR Fusion, SSL, and Efficiency", level=2)

cap4 = doc.add_paragraph()
cap4.add_run(
    "Table 4: Consolidated ablation, pillars, and efficiency.\n"
    "(a) Ablation on EfficientNet-B0, 3 seeds.  (b) SAR fusion and SSL, single seed.  (c) Efficiency on RTX 4070, batch 32."
).bold = True
cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER

t4 = doc.add_table(rows=1, cols=5)
t4.style = 'Table Grid'
for cell, txt in zip(t4.rows[0].cells, ["Section / Model", "Config", "OA (%)", "F1", "mIoU"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True

t4_rows = [
    ("(a) Ablation", "Full method (proposed)", "95.6 ± 1.9", "0.958 ± 0.018", "0.901 ± 0.039"),
    ("",             "Without FPN",            "95.3 ± 1.4", "0.955 ± 0.013", "0.893 ± 0.029"),
    ("",             "CE loss only",           "96.0 ± 1.5†","0.961 ± 0.014", "0.908 ± 0.033"),
    ("(b) Fusion",   "Optical-only",           "96.7",        "0.969",         "0.922"),
    ("",             "Optical+SAR",            "87.9",        "0.870",         "0.718"),
    ("(b) SSL",      "ImageNet init",          "96.9",        "0.969",         "0.922"),
    ("",             "SimCLR",                 "93.2",        "0.931",         "0.835"),
]
for row_data in t4_rows:
    row = t4.add_row().cells
    for cell, val in zip(row, row_data):
        cell.paragraphs[0].add_run(val).font.size = Pt(10)

# Efficiency sub-table
eff_cap = doc.add_paragraph()
eff_cap.add_run("(c) Efficiency:").bold = True

t4e = doc.add_table(rows=1, cols=4)
t4e.style = 'Table Grid'
for cell, txt in zip(t4e.rows[0].cells, ["Model", "Params/Lat. (M/ms)", "Throughput (p/s)", "GPU (MB)"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True
for row_data in [
    ("MobileNetV3-S",  "3.39 / 5.42",  "184.4", "30.5"),
    ("EfficientNet-B0","5.25 / 7.14",  "140.1", "52.9"),
    ("ResNet50",       "11.31 / 5.02", "199.2", "83.4"),
    ("Swin-Tiny",      "29.83 / 10.98","91.1",  "176.3"),
]:
    row = t4e.add_row().cells
    for cell, val in zip(row, row_data):
        cell.paragraphs[0].add_run(val).font.size = Pt(10)

doc.add_paragraph()
ablation_items = [
    ("Ablation —", "FPN adds +0.3 OA points, dominated by Transition-class boundaries. CE-only's 0.4-point OA advantage over the full method is not statistically significant (p=0.71); per-class evaluation shows the full method achieves Transition F1 = 0.98 — the minority encroachment-relevant class — justifying the combined loss despite its marginal aggregate OA trade-off."),
    ("SAR Fusion —", "SAR–optical fusion underperforms optical-only because only 33% of patches have paired SAR and the available SAR is post-monsoon vs. pre-monsoon optical; this is a deliberately reported negative result since published studies show SAR only helps under strict temporal alignment."),
    ("SSL —", "At our label budget (2,730 patches), ImageNet transfer beats SimCLR, whose advantage emerges mainly under <500 labels."),
    ("Efficiency —", "ResNet50 offers the best accuracy–latency balance (5.02 ms, 199 patches/s, 83 MB); MobileNetV3-Small is the edge-device model at 30.5 MB; Swin-Tiny is heaviest but most transferable."),
]
for title, detail in ablation_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title + " ")
    r1.italic = True
    p.add_run(detail)

add_heading("D. Urban Growth Forecasting and Encroachment Alerting", level=2)
forecast_items = [
    ("How accurately can we predict future sprawl?",
     "The Bi-LSTM, trained on Census/RBI socio-economic features and satellite-derived urban areas, reaches R² = 0.9564 "
     "(MAE = 119.5 km², MAPE = 6.66%). With only four real satellite anchor years, R² drops to 0.559 vs. Ridge (R² = 0.601). "
     "The LSTM's key advantage is uncertainty quantification via MC Dropout: the 95% CI bands quantify infrastructure "
     "investment risk — a capability absent from every published Indian forecasting study."),
    ("Can encroachment be automatically flagged?",
     "On the LEVIR-CD benchmark, our Siamese EfficientNet-B0 reaches F1 = 0.949 (OA = 94.5%). The three-head change "
     "detector reaches 99.33% binary change accuracy at 77.5 ms latency; a seven-city simulation produces 55 simulated "
     "alerts (4 CRITICAL, 28 HIGH, 2 MED, 21 LOW) and 11 protected-zone violations, with critical alerts near Sanjay "
     "Gandhi NP and Pallikaranai marsh routed to MoEFCC."),
    ("How does accuracy vary across cities?",
     "Mumbai achieves 95.1%, Delhi NCR 94.4%, and Bangalore 80.5%. Bangalore's Non-Urban recall is only 10.5% because "
     "dispersed IT-park vegetation fragments resemble Transition zones; city-wise feature clustering corroborates the LOCO drop."),
    ("Does the model generalize across time?",
     "Applied without retraining to 2019 and 2023 Sentinel-2, the 2021-trained ResNet50 estimates Mumbai's urban extent "
     "at 1,161 and 1,451 km² (+25%, consistent with coastal construction), while Delhi NCR and Bangalore remain saturated "
     "inside their tight bounding boxes."),
]
for title, detail in forecast_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + " ")
    r1.italic = True
    p.add_run(detail)

# ═══════════════════════════════════════════════════════════════════════════════
#  VI. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("VI. Conclusion", level=1)
add_para(
    "We presented an end-to-end Indian urban expansion framework unifying classification, fusion, self-supervision, "
    "Bi-LSTM forecasting, and regulatory alerting. On real Mumbai/Delhi NCR/Bangalore data, ResNet50 reaches "
    "97.5 ± 0.2% OA — six points over the best published Indian SVM. The first Indian LOCO benchmark reveals a "
    "15–20% cross-city gap and a CNN–Transformer ranking reversal (Swin-Tiny generalizes best). The pipeline produces "
    "uncertainty-calibrated 2024–2035 forecasts and 55 simulated encroachment alerts with authority routing. "
    "The study is currently limited to three Tier-1 cities, WorldCover weak labels, a dilated Transition class, "
    "calibrated socio-economic forecasting inputs, simulation-based alerts, and seasonally mismatched SAR. "
    "Ethical considerations: alerts are advisory; human-in-the-loop verification is required to avoid disparate "
    "impact on informal settlements."
)

fw = doc.add_paragraph()
fw.paragraph_format.space_before = Pt(4)
r = fw.add_run("Future work: ")
r.italic = True
fw.add_run(
    "We plan to expand to Tier-2/3 cities, validate alerts against ground-truth municipal records, "
    "and improve SAR fusion via strict temporal alignment."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  ACKNOWLEDGMENT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Acknowledgment", level=1)
add_para(
    "The authors thank PES University, Bengaluru, for providing the research infrastructure and academic environment "
    "to conduct this work. We also thank ESA for Sentinel imagery (Copernicus), USGS for Landsat data, Google for the "
    "Earth Engine platform, and MoEFCC for regulatory-zone documentation."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Appendix I: Reproducibility Protocol", level=1)
repro_items = [
    ("Tools and Hardware:",
     "GEE Python API (v0.1.374) for satellite data export (~5.7 GB); Python 3.11, PyTorch 2.0.1+cu118, "
     "torchvision 0.15.2, timm 0.9.12, scikit-learn 1.3.2; NVIDIA RTX 4070 Laptop GPU (8 GB VRAM), 16 GB RAM, Windows 11."),
    ("Seeds:",
     "{42, 123, 7}; CUDA determinism enabled; all results report mean ± std."),
    ("Hyperparameters:",
     "Batch 32, 256×256 patches, 6 S2 bands; LR 10⁻³/10⁻⁴/10⁻⁵ (stages 1/2/3), cosine annealing, patience 10; "
     "loss CE 0.6 + Focal 0.3 (γ=2) + Dice 0.1, class weights [1,1,3]; augmentation H/VFlip, RandRot, ColorJitter, "
     "Mixup α=0.2. SVM: RBF, C∈{1,10,100}, PCA(99%), NDVI/NDBI/NDWI/SAVI/BSI + GLCM. "
     "Bi-LSTM: hidden 64, 1 layer, 2-head attention, MC Dropout 0.2, Huber loss, 50 epochs. SimCLR: τ=0.07, 100 epochs."),
    ("Data/Code:",
     "GEE project urban-expansion-india; code and data at https://github.com/kaustubhagarwal21/urban-expansion-monitoring (upon acceptance)."),
]
for title, detail in repro_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + " ")
    r1.italic = True
    p.add_run(detail)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"c:\Users\KAUSTUBH\Desktop\AISD PROJECT\paper\main.docx"
doc.save(out)
print(f"Saved: {out}")
